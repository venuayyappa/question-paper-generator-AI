from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from io import BytesIO
from datetime import datetime
import re

def build_docx(subject_name: str,
               course_code: str,
               exam_type: str,
               semester: str,
               total_marks: int,
               question_paper: str,
               answer_key: str,
               instructions: str = "") -> BytesIO:
    """
    Creates an in-memory .docx file with STRICT VTU / Autonomous formatting.
    Returns a BytesIO object.
    """
    doc = Document()
    
    # helper for styling
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # --- 1. HEADER SECTION (No Duplication) ---
    # Institute Name
    p = doc.add_paragraph("INSTITUTE OF ENGINEERING AND TECHNOLOGY")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)

    # Exam Title
    p = doc.add_paragraph(f"{exam_type.upper()} EXAMINATION")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    
    # Helper to add a line with left/right content using tabs
    def add_header_line(left_text, right_text):
        p = doc.add_paragraph()
        p.paragraph_format.tab_stops.clear_all()
        # Add a tab stop at the far right margin (approx 6.5 inches)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
        
        run = p.add_run(f"{left_text}\t{right_text}")
        run.bold = True

    # Program and Semester
    add_header_line(f"Program: B.E.", f"Semester: {semester}")
    
    # Course Name and Max Marks
    add_header_line(f"Course Name: {subject_name}", f"Max. Marks: {total_marks}")
    
    # Course Code and Duration
    add_header_line(f"Course Code: {course_code}", "Duration: 3 Hrs")

    doc.add_paragraph("")  # Spacer line

    # --- 2. INSTRUCTIONS ---
    if instructions:
        p = doc.add_paragraph("Instructions to the Candidates:")
        p.runs[0].bold = True
        
        # Clean and add bullets
        lines = instructions.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            # Remove leading "1.", "1)", "-", "*" etc.
            clean_line = re.sub(r'^[\d\.\)\-\*]+\s*', '', line)
            doc.add_paragraph(clean_line, style='List Bullet')
            
    doc.add_paragraph("") # Spacer

    # --- 3. QUESTION PAPER BODY ---
    
    # Metadata keys to strip from Gemini output
    metadata_keys = ["SUBJECT:", "COURSE CODE:", "EXAM TYPE:", "SEMESTER:", "TOTAL MARKS:", "TIME:", "INSTRUCTIONS:", "NOTE:", "PROGRAM:"]

    def clean_text(text_line):
        # Remove markdown chars
        text_line = text_line.replace('**', '').replace('__', '').replace('###', '')
        # Remove "QUESTION PAPER" heading if Gemini outputs it
        if text_line.upper() == "QUESTION PAPER":
            return ""
        return text_line.strip()

    def is_metadata(text_line):
        upper_line = text_line.upper()
        for key in metadata_keys:
            if upper_line.startswith(key):
                return True
        return False

    def is_separator(text_line):
        # Checks for lines like "---" or "***"
        if not text_line: return False
        unique = set(text_line)
        return unique.issubset({'-', '_', '*', ' '})

    for line in question_paper.splitlines():
        line = clean_text(line)
        if not line:
            continue
            
        if is_metadata(line):
            continue
            
        if is_separator(line):
            continue

        upper_line = line.upper()

        # UNIT / SECTION Detection -> Bold Centered
        if upper_line.startswith("UNIT") or upper_line.startswith("SECTION") or upper_line.startswith("PART"):
            # Ensure format is "UNIT - I" etc
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.bold = True
            p.paragraph_format.space_before = Pt(12)

        # CO / Marks Line Detection (e.g. "CO1 (10)")
        # Regex: Look for CO followed by digits, optional space, parenthesized number
        elif re.search(r'CO\d+\s*\(\d+\)', line):
            # Right align this line
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Sub-questions (a), b), i., ii.) -> Indent
        elif re.match(r'^([a-z]\)|[ivx]+\.)', line):
            p = doc.add_paragraph(line)
            p.paragraph_format.left_indent = Inches(0.5)

        # Normal Text / Questions
        else:
            doc.add_paragraph(line)

    # --- 4. ANSWER KEY ---
    doc.add_page_break()
    heading = doc.add_heading("Answer Key", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for line in answer_key.splitlines():
        line = clean_text(line)
        if not line: continue
        if is_metadata(line): continue
        if is_separator(line): continue
        
        doc.add_paragraph(line)

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
